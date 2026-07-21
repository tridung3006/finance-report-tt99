"""Generate the checked-in TT99 template schema from the official DOCX files.

The output is deterministic and intentionally contains only document structure/text;
it never reads or writes accounting data.
"""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_DIR = ROOT / "templates" / "tt99"
OUTPUT = ROOT / "server" / "tt99-template-schema.json"
MAJOR_HEADING = re.compile(r"^(I|II|III|IV|V|VI|VII|VIII|IX|X)\.\s")
PLACEHOLDER = re.compile(r"^[.\u2026\s]+$")


def clean_text(value: str) -> str:
    return " ".join(value.replace("\xa0", " ").split())


def iter_body_blocks(document: Document):
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield "paragraph", Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield "table", Table(child, document)


def report_lines(path: Path, code_pattern: str) -> list[dict[str, str]]:
    pattern = re.compile(code_pattern)
    found: list[dict[str, str]] = []
    seen: set[str] = set()
    for table in Document(path).tables:
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            for index, value in enumerate(cells):
                if not pattern.fullmatch(value) or value in seen:
                    continue
                label = cells[index - 1] if index > 0 else ""
                if not label and index + 1 < len(cells):
                    label = cells[index + 1]
                found.append({"code": value, "label": label})
                seen.add(value)
                break
    return found


def cell_value(value: str):
    text = clean_text(value)
    return None if text and PLACEHOLDER.fullmatch(text) else text


def b09_schema(path: Path) -> dict:
    document = Document(path)
    sections: list[dict] = []
    current = None
    content_table_index = 0
    word_table_index = -1

    for kind, block in iter_body_blocks(document):
        if kind == "paragraph":
            text = clean_text(block.text)
            if not text:
                continue
            if MAJOR_HEADING.match(text):
                current = {"title": text, "blocks": []}
                sections.append(current)
            elif current is not None:
                current["blocks"].append({"type": "paragraph", "text": text})
            continue

        word_table_index += 1
        if word_table_index == 0 or word_table_index == len(document.tables) - 1:
            continue
        if current is None:
            raise RuntimeError(f"B09 content table {word_table_index} appears before section I")

        content_table_index += 1
        rows = [[cell_value(cell.text) for cell in row.cells] for row in block.rows]
        column_count = max((len(row) for row in rows), default=0)
        if any(len(row) != column_count for row in rows):
            raise RuntimeError(f"B09 table {content_table_index} has an inconsistent grid")

        title = ""
        if current["blocks"] and current["blocks"][-1]["type"] == "paragraph":
            title = current["blocks"].pop()["text"]
        current["blocks"].append({
            "type": "table",
            "table": {
                "templateIndex": content_table_index,
                "title": title,
                "columnCount": column_count,
                "rowCount": len(rows),
                "rows": rows,
            },
        })

    tables = [
        block["table"]
        for section in sections
        for block in section["blocks"]
        if block["type"] == "table"
    ]
    if len(tables) != 53:
        raise RuntimeError(f"Expected 53 B09 content tables, found {len(tables)}")

    return {
        "wordTableCount": len(document.tables),
        "contentTableCount": len(tables),
        "sections": sections,
    }


def main() -> None:
    schema = {
        "source": "Thông tư 99/2025/TT-BTC - DOCX supplied by the user",
        "B01": report_lines(TEMPLATE_DIR / "b01-dn.docx", r"\d{3}[ab]?"),
        "B03": report_lines(TEMPLATE_DIR / "b03-dn-truc-tiep.docx", r"\d{2}"),
        "B09": b09_schema(TEMPLATE_DIR / "b09-dn.docx"),
    }
    OUTPUT.write_text(json.dumps(schema, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(
        f"Wrote {OUTPUT.relative_to(ROOT)}: "
        f"B01={len(schema['B01'])}, B03={len(schema['B03'])}, "
        f"B09 tables={schema['B09']['contentTableCount']}"
    )


if __name__ == "__main__":
    main()
